import os
import PyPDF2
import docx
import magic
import hashlib
from datetime import datetime
from werkzeug.utils import secure_filename
from backend.models.document import Document, db
import json

class DocumentProcessor:
    def extract_text_from_txt(self, filepath):
        """Extract text content from plain text file"""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                text = file.read()
            return {
                'content': text.strip(),
                'page_count': 1,
                'word_count': len(text.split())
            }
        except Exception as e:
            raise Exception(f"Error extracting text file content: {str(e)}")
    def __init__(self, upload_folder='uploads', processed_folder='processed_docs'):
        self.upload_folder = upload_folder
        self.processed_folder = processed_folder
        self.allowed_extensions = {'pdf', 'docx', 'doc', 'txt', 'ppt', 'pptx', 'xls', 'xlsx'}
        # Ensure directories exist
        os.makedirs(upload_folder, exist_ok=True)
        os.makedirs(processed_folder, exist_ok=True)
    
    def is_allowed_file(self, filename):
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in self.allowed_extensions
    
    def get_file_type(self, filepath):
        """Detect file type using python-magic"""
        try:
            mime = magic.from_file(filepath, mime=True)
            return mime
        except:
            # Fallback to extension-based detection
            ext = filepath.rsplit('.', 1)[1].lower()
            if ext == 'pdf':
                return 'application/pdf'
            elif ext in ['docx', 'doc']:
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            return 'unknown'
    
    def generate_unique_filename(self, original_filename):
        """Generate unique filename using timestamp and hash"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        file_hash = hashlib.md5(original_filename.encode()).hexdigest()[:8]
        name, ext = os.path.splitext(original_filename)
        return f"{timestamp}_{file_hash}_{secure_filename(name)}{ext}"
    
    def save_uploaded_file(self, file):
        """Save uploaded file and return file info"""
        if not file or not self.is_allowed_file(file.filename):
            raise ValueError("Invalid file type")
        
        original_filename = file.filename
        unique_filename = self.generate_unique_filename(original_filename)
        filepath = os.path.join(self.upload_folder, unique_filename)
        
        file.save(filepath)
        
        # Get file info
        file_size = os.path.getsize(filepath)
        file_type = self.get_file_type(filepath)
        
        return {
            'filepath': filepath,
            'filename': unique_filename,
            'original_filename': original_filename,
            'file_size': file_size,
            'file_type': file_type
        }
    
    def extract_text_from_pdf(self, filepath):
        """Extract text content from PDF file"""
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                page_count = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    'content': text.strip(),
                    'page_count': page_count,
                    'word_count': len(text.split())
                }
        except Exception as e:
            raise Exception(f"Error extracting PDF content: {str(e)}")
    
    def extract_text_from_docx(self, filepath):
        """Extract text content from Word document"""
        try:
            doc = docx.Document(filepath)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return {
                'content': text.strip(),
                'page_count': 1,  # Word docs don't have clear page breaks in python-docx
                'word_count': len(text.split())
            }
        except Exception as e:
            raise Exception(f"Error extracting Word document content: {str(e)}")
    
    def extract_content(self, filepath, file_type):
        """Extract content based on file type, with extension fallback"""
        ft = file_type.lower()
        ext = os.path.splitext(filepath)[1].lower()
        if 'pdf' in ft or ext == '.pdf':
            return self.extract_text_from_pdf(filepath)
        elif 'word' in ft or 'document' in ft or ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(filepath)
        elif 'text' in ft or 'plain' in ft or ext == '.txt':
            return self.extract_text_from_txt(filepath)
        elif 'powerpoint' in ft or ext in ['.ppt', '.pptx']:
            return {
                'content': '',
                'page_count': 0,
                'word_count': 0
            }  # Placeholder for PowerPoint
        elif 'excel' in ft or ext in ['.xls', '.xlsx']:
            return {
                'content': '',
                'page_count': 0,
                'word_count': 0
            }  # Placeholder for Excel
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def generate_title(self, content, filename):
        """Generate document title from content or filename"""
        if content:
            # Try to extract title from first few lines
            lines = content.split('\n')[:5]
            for line in lines:
                line = line.strip()
                if line and len(line) > 10 and len(line) < 100:
                    return line
        
        # Fallback to filename without extension
        return os.path.splitext(filename)[0].replace('_', ' ').title()
    
    def generate_summary(self, content, max_length=500):
        """Generate document summary from content"""
        if not content:
            return ""
        
        # Simple extractive summary - take first few sentences
        sentences = content.split('.')[:3]
        summary = '. '.join(sentences).strip()
        
        if len(summary) > max_length:
            summary = summary[:max_length] + "..."
        
        return summary
    
    def process_document(self, file):
        """Complete document processing pipeline"""
        try:
            # Save uploaded file
            file_info = self.save_uploaded_file(file)
            
            # Create document record
            document = Document(
                filename=file_info['filename'],
                original_filename=file_info['original_filename'],
                file_type=file_info['file_type'],
                file_size=file_info['file_size'],
                processing_status='processing'
            )
            
            db.session.add(document)
            db.session.commit()
            
            try:
                # Extract content
                content_info = self.extract_content(file_info['filepath'], file_info['file_type'])
                
                # Generate title and summary
                title = self.generate_title(content_info['content'], file_info['original_filename'])
                summary = self.generate_summary(content_info['content'])
                
                # Update document with extracted content
                document.content = content_info['content']
                document.title = title
                document.summary = summary
                document.word_count = content_info['word_count']
                document.page_count = content_info['page_count']
                document.processed_date = datetime.utcnow()
                document.processing_status = 'completed'
                
                # Save processed content to file
                processed_filepath = os.path.join(
                    self.processed_folder, 
                    f"{document.id}_{file_info['filename']}.txt"
                )
                with open(processed_filepath, 'w', encoding='utf-8') as f:
                    f.write(content_info['content'])
                
                db.session.commit()
                
                return document
                
            except Exception as e:
                document.processing_status = 'failed'
                document.error_message = str(e)
                db.session.commit()
                raise e
                
        except Exception as e:
            if 'document' in locals():
                db.session.rollback()
            raise e